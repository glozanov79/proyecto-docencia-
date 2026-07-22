"""
Genera la guía de videos (.docx) a partir de un brief.md, buscando videos reales
de YouTube con la herramienta de búsqueda web de Claude.

Requiere:
- Variable de entorno ANTHROPIC_API_KEY

Uso:
    python generar_videos.py salidas/2026-08-05/brief.md
"""

import json
import os
import re
import sys
from pathlib import Path

import anthropic
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

sys.stdout.reconfigure(encoding="utf-8")
sys.stderr.reconfigure(encoding="utf-8")

MODEL = "claude-sonnet-4-6"

SYSTEM_PROMPT = """Eres un curador de recursos educativos en YouTube para un curso
universitario. Vas a recibir el brief de una clase y debes buscar, con la herramienta
de búsqueda web, entre 4 y 5 videos REALES de YouTube (deben existir de verdad,
nunca inventados) que expliquen los subtemas de "Contenido clave a cubrir" del brief,
coherentes con la materia del curso descrito en el brief.

REGLAS:
- Usa la búsqueda web para confirmar que cada video existe antes de incluirlo.
- Cubre distintos subtemas del brief, no solo uno.
- Para cada video incluye: título exacto, URL completa de YouTube, y una frase de
  por qué sirve para este tema.
- Al final de tu respuesta, y SOLO al final, responde ÚNICAMENTE con un bloque JSON
  (sin markdown, sin texto después) con esta forma exacta:
  {"videos": [{"titulo": "...", "url": "...", "por_que": "..."}]}
"""


def extraer_json(texto):
    texto = texto.strip()
    if texto.startswith("```"):
        texto = re.sub(r"^```(json)?", "", texto).rstrip("`").strip()
    inicio = texto.find("{")
    fin = texto.rfind("}")
    return json.loads(texto[inicio:fin + 1])


def generar_videos(brief_texto, api_key):
    client = anthropic.Anthropic(api_key=api_key)
    respuesta = client.messages.create(
        model=MODEL,
        max_tokens=4000,
        system=SYSTEM_PROMPT,
        tools=[{"type": "web_search_20250305", "name": "web_search"}],
        messages=[
            {"role": "user", "content": f"Brief de la clase:\n\n{brief_texto}\n\nBusca los videos y responde con el JSON."}
        ],
    )
    texto = "".join(b.text for b in respuesta.content if b.type == "text")
    return extraer_json(texto)


def agregar_hipervinculo(paragraph, url, texto):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = texto
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def escribir_docx(datos, ruta_salida):
    doc = Document()
    doc.add_heading("Guía de videos sugeridos", level=1)

    for video in datos.get("videos", []):
        p_titulo = doc.add_paragraph()
        run = p_titulo.add_run(video.get("titulo", "(sin título)"))
        run.bold = True
        run.font.size = Pt(12)

        p_link = doc.add_paragraph()
        agregar_hipervinculo(p_link, video.get("url", ""), video.get("url", ""))

        doc.add_paragraph(video.get("por_que", ""))
        doc.add_paragraph()

    doc.save(ruta_salida)


def main():
    if len(sys.argv) != 2:
        print("Uso: python generar_videos.py <ruta_al_brief.md>")
        sys.exit(1)

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("ERROR: define la variable de entorno ANTHROPIC_API_KEY antes de correr esto.")
        sys.exit(1)

    brief_path = Path(sys.argv[1])
    brief_texto = brief_path.read_text(encoding="utf-8")
    salida_dir = brief_path.parent

    print("Llamando a la API de Claude (con búsqueda web) para encontrar videos...")
    datos = generar_videos(brief_texto, api_key)

    ruta_salida = salida_dir / "videos.docx"
    escribir_docx(datos, ruta_salida)
    print(f"videos.docx generado en: {ruta_salida}")


if __name__ == "__main__":
    main()
